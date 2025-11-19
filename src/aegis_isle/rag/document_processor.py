"""
Document processing components for RAG pipeline.
"""

import hashlib
import mimetypes
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from pydantic import BaseModel, Field

from ..core.logging import logger


class DocumentMetadata(BaseModel):
    """Metadata for a document."""

    filename: str
    file_path: Optional[str] = None
    file_size: int = 0
    mime_type: Optional[str] = None
    language: Optional[str] = "en"
    source: str = "upload"  # upload, url, database, etc.
    author: Optional[str] = None
    title: Optional[str] = None
    created_at: datetime = Field(default_factory=datetime.now)
    modified_at: Optional[datetime] = None
    tags: List[str] = Field(default_factory=list)
    custom_fields: Dict[str, Any] = Field(default_factory=dict)


class DocumentChunk(BaseModel):
    """Represents a chunk of a document."""

    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    content: str
    chunk_index: int
    start_pos: int = 0
    end_pos: int = 0
    metadata: Dict[str, Any] = Field(default_factory=dict)
    embedding: Optional[List[float]] = None
    created_at: datetime = Field(default_factory=datetime.now)


class ProcessedDocument(BaseModel):
    """Represents a processed document."""

    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    content: str
    metadata: DocumentMetadata
    chunks: List[DocumentChunk] = Field(default_factory=list)
    processing_stats: Dict[str, Any] = Field(default_factory=dict)
    content_hash: Optional[str] = None
    created_at: datetime = Field(default_factory=datetime.now)

    def __post_init__(self):
        """Calculate content hash after initialization."""
        if not self.content_hash and self.content:
            self.content_hash = hashlib.sha256(
                self.content.encode('utf-8')
            ).hexdigest()


class DocumentProcessor:
    """Main document processor that handles different file types."""

    def __init__(self):
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_doc,
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.html': self._process_html,
            '.htm': self._process_html,
        }
        self.ocr_enabled = True

    async def process_file(
        self,
        file_path: Union[str, Path],
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """Process a file and return a ProcessedDocument."""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        logger.info(f"Processing file: {file_path}")

        # Determine file type
        file_ext = file_path.suffix.lower()
        if file_ext not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_ext}")

        # Create metadata
        doc_metadata = self._create_metadata(file_path, metadata or {})

        # Process content based on file type
        processor = self.supported_types[file_ext]
        content, processing_stats = await processor(file_path)

        # Create processed document
        document = ProcessedDocument(
            content=content,
            metadata=doc_metadata,
            processing_stats=processing_stats
        )

        logger.info(
            f"Processed document {document.id}: "
            f"{len(content)} characters, "
            f"{processing_stats.get('processing_time', 0):.2f}s"
        )

        return document

    async def process_text(
        self,
        content: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """Process raw text content."""
        doc_metadata = DocumentMetadata(
            filename="text_input",
            mime_type="text/plain",
            source="text_input",
            **(metadata or {})
        )

        document = ProcessedDocument(
            content=content,
            metadata=doc_metadata,
            processing_stats={"content_length": len(content)}
        )

        return document

    async def process_url(
        self,
        url: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """Process content from a URL."""
        import aiohttp

        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url) as response:
                    content = await response.text()

            doc_metadata = DocumentMetadata(
                filename=url.split('/')[-1] or "webpage",
                source="url",
                mime_type=response.headers.get('content-type', 'text/html'),
                **(metadata or {}),
                custom_fields={"url": url}
            )

            # Process HTML content
            if 'html' in doc_metadata.mime_type:
                content, stats = await self._process_html_content(content)
            else:
                stats = {"content_length": len(content)}

            document = ProcessedDocument(
                content=content,
                metadata=doc_metadata,
                processing_stats=stats
            )

            return document

        except Exception as e:
            logger.error(f"Error processing URL {url}: {e}")
            raise

    def _create_metadata(
        self,
        file_path: Path,
        custom_metadata: Dict[str, Any]
    ) -> DocumentMetadata:
        """Create metadata for a file."""
        stat = file_path.stat()
        mime_type, _ = mimetypes.guess_type(str(file_path))

        return DocumentMetadata(
            filename=file_path.name,
            file_path=str(file_path),
            file_size=stat.st_size,
            mime_type=mime_type,
            created_at=datetime.fromtimestamp(stat.st_ctime),
            modified_at=datetime.fromtimestamp(stat.st_mtime),
            **custom_metadata
        )

    async def _process_pdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process PDF file."""
        import time
        start_time = time.time()

        try:
            import PyPDF2

            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                content_parts = []

                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        content_parts.append(text)

                content = "\n\n".join(content_parts)

                stats = {
                    "total_pages": len(reader.pages),
                    "processing_time": time.time() - start_time,
                    "extraction_method": "PyPDF2"
                }

                return content, stats

        except Exception as e:
            logger.warning(f"PyPDF2 failed for {file_path}: {e}")

            if self.ocr_enabled:
                return await self._process_pdf_with_ocr(file_path)
            else:
                raise

    async def _process_pdf_with_ocr(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process PDF with OCR as fallback."""
        import time
        start_time = time.time()

        try:
            import pdf2image
            import pytesseract
            from PIL import Image

            # Convert PDF to images
            images = pdf2image.convert_from_path(str(file_path))
            content_parts = []

            for page_num, image in enumerate(images):
                # Extract text using OCR
                text = pytesseract.image_to_string(image, lang='eng+chi_sim')
                if text.strip():
                    content_parts.append(text)

            content = "\n\n".join(content_parts)

            stats = {
                "total_pages": len(images),
                "processing_time": time.time() - start_time,
                "extraction_method": "OCR"
            }

            return content, stats

        except Exception as e:
            logger.error(f"OCR processing failed for {file_path}: {e}")
            raise

    async def _process_docx(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process DOCX file."""
        import time
        start_time = time.time()

        try:
            from docx import Document

            doc = Document(str(file_path))
            content_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)

            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        content_parts.append(row_text)

            content = "\n".join(content_parts)

            stats = {
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables),
                "processing_time": time.time() - start_time
            }

            return content, stats

        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise

    async def _process_doc(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process DOC file (legacy Word format)."""
        # For .doc files, we'd typically use python-docx2txt or similar
        # For now, return an error suggesting conversion
        raise NotImplementedError(
            "Legacy .doc format not supported. Please convert to .docx"
        )

    async def _process_text(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process plain text file."""
        import time
        start_time = time.time()

        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            stats = {
                "content_length": len(content),
                "processing_time": time.time() - start_time
            }

            return content, stats

        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()

                    stats = {
                        "content_length": len(content),
                        "processing_time": time.time() - start_time,
                        "encoding": encoding
                    }

                    return content, stats
                except UnicodeDecodeError:
                    continue

            raise ValueError(f"Unable to decode text file: {file_path}")

    async def _process_markdown(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process Markdown file."""
        return await self._process_text(file_path)

    async def _process_html(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process HTML file."""
        import time
        start_time = time.time()

        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()

        content, additional_stats = await self._process_html_content(html_content)

        stats = {
            "processing_time": time.time() - start_time,
            **additional_stats
        }

        return content, stats

    async def _process_html_content(self, html_content: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from HTML content."""
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(html_content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text content
            text = soup.get_text()

            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = "\n".join(chunk for chunk in chunks if chunk)

            stats = {
                "original_length": len(html_content),
                "extracted_length": len(content),
                "extraction_method": "BeautifulSoup"
            }

            return content, stats

        except ImportError:
            logger.warning("BeautifulSoup not available, using regex fallback")
            import re

            # Simple regex-based HTML tag removal
            clean = re.compile('<.*?>')
            content = re.sub(clean, '', html_content)

            stats = {
                "original_length": len(html_content),
                "extracted_length": len(content),
                "extraction_method": "regex"
            }

            return content, stats